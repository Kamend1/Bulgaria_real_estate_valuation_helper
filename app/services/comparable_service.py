from __future__ import annotations

import io
import uuid
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from sqlalchemy import text
from sqlalchemy.dialects.postgresql import insert as pg_insert
from sqlalchemy.orm import Session

from app.db.models import AppraisalReport, ComparablePool
from app.services.analytics_service import get_market_trend
from utils.feature_engineering import PROPERTY_TYPE_DISPLAY

MAX_PINNED = 6   # max pinned-for-report per type (Word table limit)

# ── Word export constants ─────────────────────────────────────────────────────

_KNOB_TEXT = (
    "КСБ Аналитика е вписана със сертификат № 900200284 в Регистъра на Камарата на "
    "независимите оценители в България (КНОБ) и извършва оценителска дейност на "
    "територията на Република България за недвижими имоти, "
    "както и за финансови активи и финансови институции."
)
_BRAND_DARK = RGBColor(0x1E, 0x3A, 0x5F)
_BRAND_BLUE = RGBColor(0x25, 0x63, 0xEB)
_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
_MUTED      = RGBColor(0x64, 0x74, 0x8B)

# ── Purpose-differentiated front-matter boilerplate ───────────────────────────
# Real appraisal practice cites a different legal/accounting basis depending on
# what the report is FOR -- a fair-value opinion for financial reporting (IFRS
# 13 / IAS 16 / IAS 36) is legally a different thing from a valuation backing
# a non-cash capital contribution (чл. 72 ТЗ), which even ends with a clause
# neither of the other purposes need (computing a share count from the value).
# "market_opinion" is the default and matches this app's original,
# purpose-agnostic behavior -- no standard is cited unless the appraiser
# explicitly says the report is for one of the other two purposes.
_PURPOSE_TEXTS: dict[str, dict[str, str]] = {
    "market_opinion": {
        "label": "Обща пазарна консултация",
        "cel": (
            "Настоящият доклад е изготвен по конкретно възлагане с цел подпомагане на "
            "възложителя при определяне на пазарната стойност на описания имот. "
            "Докладът не е обвързан с конкретен счетоводен или нормативен стандарт и "
            "не следва да се използва за целите на финансова отчетност, съдебно "
            "производство или прехвърляне на собственост без изрично потвърждение от "
            "страна на оценителя."
        ),
        "standart": (
            "Оценката представлява становище на оценителя за пазарната стойност на "
            "имота — най-вероятната цена, по която имотът би могъл да бъде продаден на "
            "свободния пазар между желаещи купувач и продавач, при обичайни пазарни "
            "условия и без принуда от страна на нито една от страните, към датата на "
            "оценката."
        ),
    },
    "fair_value_ifrs": {
        "label": "Справедлива стойност (МСФО 13 / МСС 16 / МСС 36)",
        "cel": (
            "Целта на настоящата оценка е определяне на справедливата стойност на "
            "описания имот към датата на оценка, в съответствие с изискванията на "
            "Международен стандарт за финансово отчитане 13 (МСФО 13) и Международни "
            "счетоводни стандарти 16 и 36 (МСС 16, МСС 36), за да послужи на "
            "възложителя за целите на счетоводното отчитане на активите."
        ),
        "standart": (
            "В оценителския доклад се съблюдават изискванията на възприетите "
            "международни стандарти за финансова отчетност (МСФО 13) и международни "
            "счетоводни стандарти (МСС 16 и МСС 36). Справедливата стойност се определя "
            "като цената, която би била получена при продажба на актива или платена при "
            "прехвърлянето на задължение при обичайна сделка между пазарни участници "
            "към датата на оценката."
        ),
    },
    "noncash_contribution": {
        "label": "Непарична вноска (чл. 72, ал. 2 ТЗ)",
        "cel": (
            "Целта на настоящото становище е определяне на пазарната стойност на "
            "описания недвижим имот, който следва да послужи като непарична вноска в "
            "капитала на дружество, в съответствие с изискванията на чл. 72, ал. 2 от "
            "Търговския закон и чл. 123 от Наредба № 1 от 14.02.2007 г. за водене, "
            "съхраняване и достъп до търговския регистър."
        ),
        "standart": (
            "Оценката е изготвена съгласно приложимите Български стандарти за "
            "оценяване, приети от Камарата на независимите оценители в България. "
            "Определената стойност представлява пазарната стойност на имота към датата "
            "на оценка и служи единствено за целите на удостоверяване на съответствието "
            "на непаричната вноска с размера на записания дружествен дял, съгласно чл. "
            "72, ал. 2 от Търговския закон."
        ),
    },
}

def get_purpose_options() -> list[tuple[str, str]]:
    """[(slug, display_label), ...] for the report_purpose <select> --
    "market_opinion" first since it's the default."""
    order = ["market_opinion", "fair_value_ifrs", "noncash_contribution"]
    return [(slug, _PURPOSE_TEXTS[slug]["label"]) for slug in order]


_LIMITING_CONDITIONS_TEXT = (
    "Настоящата оценка представлява становище на оценителя, изготвено с "
    "необходимата грижа и въз основа на предоставената и обществено достъпна "
    "информация към датата на оценка. Не е извършена независима проверка на "
    "правния статут на имота — правото на собственост се приема за валидно и "
    "необременено, освен ако изрично не е посочено друго. Предоставената от "
    "възложителя информация се приема за достоверна, без тя да е била "
    "самостоятелно верифицирана. Не е извършван оглед за скрити дефекти, "
    "конструктивни проблеми или замърсяване на почвата — оценката не "
    "представлява техническа или екологична експертиза. Оценителят не носи "
    "отговорност за промени в пазарните условия, настъпили след датата на "
    "оценка. Докладът е изготвен за целта, посочена по-горе, и не следва да се "
    "използва за друга цел без изричното писмено съгласие на оценителя. "
    "Оценителят и свързаните с него лица нямат настоящ или бъдещ имуществен "
    "интерес към оценявания имот, а възнаграждението за изготвяне на доклада "
    "не е обвързано с заключената стойност."
)

# Column widths (cm) for the 11-column comparables table — total = 16.5 cm (A4 portrait)
_COMP_COL_WIDTHS = [0.7, 3.8, 1.2, 1.6, 1.6, 1.4, 0.9, 1.0, 1.0, 1.5, 1.8]

# ── Word helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _fill_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    pt: float = 9,
    color: RGBColor | None = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.size = Pt(pt)
    if color:
        run.font.color.rgb = color


def _add_field(para, field: str, pt: float = 8) -> None:
    """Append a Word field code (PAGE / NUMPAGES) as a run in para."""
    run = para.add_run()
    run.font.size = Pt(pt)
    for kind, instr in [
        ("begin", None),
        ("instrText", field),
        ("separate", None),
        ("end", None),
    ]:
        if kind == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = f" {instr} "
            run._r.append(el)
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
            run._r.append(el)


def _para_border(para, position: str = "bottom", color: str = "1E3A5F", sz: int = 6) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(sz))
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)
    pBdr.append(border)


def _set_table_fixed(tbl, widths_cm: list[float]) -> None:
    """Force fixed layout and set column widths via tblGrid."""
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    # Total width
    total = sum(int(Cm(w) / 635) for w in widths_cm)
    for tag in [qn("w:tblW")]:
        for el in tblPr.findall(tag):
            tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    # tblGrid
    existing_grid = tbl._tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl._tbl.remove(existing_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(Cm(w) / 635)))
        tblGrid.append(gc)
    # Insert tblGrid right after tblPr
    tblPr_idx = list(tbl._tbl).index(tblPr)
    tbl._tbl.insert(tblPr_idx + 1, tblGrid)


def _fmt(v, decimals: int = 0) -> str:
    if v is None:
        return "—"
    try:
        n = float(v)
        if decimals == 0:
            return f"{int(round(n)):,}".replace(",", " ")
        return f"{n:,.{decimals}f}".replace(",", " ")
    except Exception:
        return str(v)


def _floor_str_docx(item: dict) -> str:
    f, t = item.get("floor_model"), item.get("total_floors_model")
    if f is None and t is None:
        return "—"
    return f"{f if f is not None else '?'}/{t}" if t else (str(f) if f else "—")


def _write_comp_table(
    doc: Document,
    pinned: list[dict],
    ctype: str,
    report: AppraisalReport,
    stats: dict | None,
) -> None:
    ppsqm_lbl = "EUR/кв.м" if ctype == "sale" else "EUR/кв.м/мес"
    price_lbl = "Цена (EUR)" if ctype == "sale" else "Наем/мес (EUR)"

    headers = [
        "№", "Местоположение", "Площ\nкв.м", price_lbl,
        f"Цена/кв.м\n{ppsqm_lbl}", "Строит.\nтип", "Год.", "Ет./\nОбш.",
        "Кор.\n%", f"Кор.\n{ppsqm_lbl}", "Бележки",
    ]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    _set_table_fixed(tbl, _COMP_COL_WIDTHS)

    # ── Header row ────────────────────────────────────────────────
    hdr = tbl.rows[0]
    for cell, hdr_text in zip(hdr.cells, headers):
        _fill_cell(cell, hdr_text, bold=True, pt=8, color=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(cell, "1E3A5F")

    # ── Subject row (yellow) ──────────────────────────────────────
    if report.subject_address or report.subject_city:
        srow = tbl.add_row()
        subj_vals = [
            "Обект",
            f"{report.subject_address or ''} {report.subject_city or ''}".strip(),
            _fmt(report.subject_area_sqm), "", "",
            report.subject_construction or "—",
            str(report.subject_year) if report.subject_year else "—",
            f"{report.subject_floor or '?'}/{report.subject_total_floors or '?'}",
            "", "", "",
        ]
        for cell, val in zip(srow.cells, subj_vals):
            _fill_cell(cell, val, bold=(val == "Обект"), pt=8.5)
            _set_cell_bg(cell, "FEF3C7")

    # ── Data rows ─────────────────────────────────────────────────
    for pos, item in enumerate(pinned, start=1):
        adj = float(item.get("adjustment_pct") or 0)
        row = tbl.add_row()
        location = (
            f"{item.get('title_city_model') or ''} "
            f"{item.get('title_geo_2_model') or ''} "
            f"{item.get('location_raw') or ''}"
        ).strip()
        vals = [
            str(pos), location,
            _fmt(item.get("area_sqm_model")),
            _fmt(item.get("total_price")),
            _fmt(item.get("price_per_sqm_model")),
            item.get("construction_type_model") or "—",
            str(item["construction_year_model"]) if item.get("construction_year_model") else "—",
            _floor_str_docx(item),
            f"{adj:+.1f}%" if adj != 0 else "—",
            _fmt(item.get("adj_ppsqm")),
            item.get("analyst_note") or "",
        ]
        for cell, val in zip(row.cells, vals):
            _fill_cell(cell, val, pt=8.5)
        if item.get("pinned_for_report"):
            for cell in row.cells:
                _set_cell_bg(cell, "DCFCE7")

    # ── Stats row (light blue) ────────────────────────────────────
    if stats:
        srow = tbl.add_row()
        stats_vals = [
            "Статистики",
            f"N={stats['n']}",
            f"{_fmt(stats['min_area'])}–{_fmt(stats['max_area'])} кв.м",
            f"Ср. {_fmt(stats['mean_price'])} EUR",
            f"Мин {_fmt(stats['min_ppsqm'])} | Ср {_fmt(stats['mean_ppsqm'])} | Макс {_fmt(stats['max_ppsqm'])}",
            f"Медиана: {_fmt(stats['median'])}",
            "", "",
            f"Q25: {_fmt(stats['p25'])}",
            f"Q75: {_fmt(stats['p75'])}",
            "",
        ]
        for cell, val in zip(srow.cells, stats_vals):
            _fill_cell(cell, val, bold=(val.startswith("Статистики")), pt=8)
            _set_cell_bg(cell, "EFF6FF")


def _write_adjustment_breakdown_table(doc: Document, pinned: list[dict]) -> None:
    """Supplementary table showing the named-factor breakdown (F4/Tier 3) for
    whichever pinned comparables use structured adjustment_factors instead of
    a single blended %. Kept separate from the main comparable table (which
    stays at its existing 11 columns, sized for A4 portrait) rather than
    adding factor columns there."""
    factored = [item for item in pinned if item.get("adjustment_factors")]
    if not factored:
        return

    doc.add_paragraph()
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Разбивка на корекциите по фактори")
    r_title.bold = True
    r_title.font.size = Pt(10)

    factor_keys = list(ADJUSTMENT_FACTOR_LABELS.keys())
    headers = ["Местоположение"] + [ADJUSTMENT_FACTOR_LABELS[k] for k in factor_keys] + ["Общо"]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for cell, text_ in zip(hdr.cells, headers):
        _fill_cell(cell, text_, bold=True, pt=8, color=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(cell, "1E3A5F")

    for item in factored:
        row = tbl.add_row()
        location = (
            f"{item.get('title_city_model') or ''} {item.get('title_geo_2_model') or ''}"
        ).strip()
        factors = item["adjustment_factors"]
        vals = [location] + [
            (f"{factors[k]:+g}%" if k in factors else "—") for k in factor_keys
        ] + [f"{item.get('adjustment_pct', 0):+g}%"]
        for cell, val in zip(row.cells, vals):
            _fill_cell(cell, val, pt=8)


def _market_context_paragraph(db: Session, report: AppraisalReport) -> str | None:
    """Short, data-grounded market-context sentence for the sales-approach
    section, built from the app's own scraped listings (analytics_service's
    materialized-view trend) filtered to the subject's own segment -- not
    external macro data, which the app has no source for. Returns None if
    the subject isn't classified enough yet, or there's no trend data."""
    if not report.subject_geo_category or not report.subject_property_type:
        return None

    trend = get_market_trend(
        db, deal_type="sale",
        geo_category=report.subject_geo_category,
        property_type_slug=report.subject_property_type,
        n_runs=6,
    )
    if not trend or not trend[-1].get("median_ppsqm"):
        return None

    latest = trend[-1]
    type_label = PROPERTY_TYPE_DISPLAY.get(report.subject_property_type, report.subject_property_type)
    text_out = (
        f"Пазарен контекст: към {latest['run_date']} медианната пазарна цена за сегмент "
        f"„{type_label}“, гео-категория „{report.subject_geo_category}“, е "
        f"{round(latest['median_ppsqm'])} EUR/кв.м (на база {latest['n_listings']} обяви)."
    )
    first = trend[0]
    if len(trend) >= 2 and first.get("median_ppsqm"):
        change_pct = (latest["median_ppsqm"] - first["median_ppsqm"]) / first["median_ppsqm"] * 100
        text_out += (
            f" За периода {first['run_date']} – {latest['run_date']} медианата се "
            f"измени с {change_pct:+.1f}%."
        )
    return text_out


# ── Word / DOCX report generation ────────────────────────────────────────────

def generate_docx(db: Session, report: AppraisalReport) -> io.BytesIO:
    pool_sale = get_pool_with_stats(db, "sale", report.id)
    pool_rent = get_pool_with_stats(db, "rent", report.id)
    pinned_sale = [r for r in pool_sale["rows"] if r["pinned_for_report"]]
    pinned_rent = [r for r in pool_rent["rows"] if r["pinned_for_report"]]

    doc = Document()

    # ── Page setup (A4 portrait) ──────────────────────────────────
    sec = doc.sections[0]
    sec.page_width   = Cm(21)
    sec.page_height  = Cm(29.7)
    sec.left_margin  = Cm(2.5)
    sec.right_margin = Cm(2.0)
    sec.top_margin   = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    # ── Header ────────────────────────────────────────────────────
    sec.header.is_linked_to_previous = False   # register headerReference
    hp = sec.header.paragraphs[0]              # re-access for live reference
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    logo_path = Path("static/logo.png")
    if logo_path.exists():
        hp.add_run().add_picture(str(logo_path), height=Cm(1.1))
        hp.add_run("  ")

    r_co = hp.add_run("КСБ Аналитика")
    r_co.bold = True
    r_co.font.size = Pt(11)
    r_co.font.color.rgb = _BRAND_DARK

    r_sep = hp.add_run("  |  ДОКЛАД ЗА ОЦЕНКА НА НЕДВИЖИМ ИМОТ")
    r_sep.font.size = Pt(9)
    r_sep.font.color.rgb = _MUTED
    _para_border(hp, "bottom", "1E3A5F", 6)

    # ── Footer ────────────────────────────────────────────────────
    sec.footer.is_linked_to_previous = False   # register footerReference
    fp = sec.footer.paragraphs[0]              # re-access for live reference
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_border(fp, "top", "1E3A5F", 4)

    r_knob = fp.add_run(_KNOB_TEXT + "   |   стр. ")
    r_knob.font.size = Pt(7.5)
    r_knob.font.color.rgb = _MUTED
    _add_field(fp, "PAGE", pt=7.5)
    r_of = fp.add_run(" / ")
    r_of.font.size = Pt(7.5)
    _add_field(fp, "NUMPAGES", pt=7.5)

    # ── Cover page ────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("КСБ Аналитика")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = _BRAND_DARK

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("ДОКЛАД ЗА ОЦЕНКА НА НЕДВИЖИМ ИМОТ")
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = _BRAND_BLUE

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(report.title or "Нов доклад")
    r3.bold = True
    r3.font.size = Pt(16)
    r3.font.color.rgb = _BRAND_DARK

    doc.add_paragraph()

    if report.subject_address or report.subject_city:
        pa = doc.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ra = pa.add_run(
            f"{report.subject_address or ''}, {report.subject_city or ''}".strip(", ")
        )
        ra.font.size = Pt(12)

    if report.valuation_date:
        pv = doc.add_paragraph()
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = pv.add_run(f"Дата на оценка: {report.valuation_date.strftime('%d.%m.%Y')}")
        rv.font.size = Pt(11)
        rv.font.color.rgb = _MUTED

    doc.add_page_break()

    # ── Въведение (unnumbered, precedes the numbered sections below) ──────
    purpose_texts = _PURPOSE_TEXTS.get(report.report_purpose, _PURPOSE_TEXTS["market_opinion"])
    appraiser = report.owner if report.owner else None

    h_intro = doc.add_heading("Въведение", level=1)
    if h_intro.runs:
        h_intro.runs[0].font.color.rgb = _BRAND_DARK

    def _intro_block(title: str, body: str) -> None:
        p_title = doc.add_paragraph()
        r_title = p_title.add_run(title)
        r_title.bold = True
        r_title.font.size = Pt(10.5)
        p_body = doc.add_paragraph(body)
        p_body.runs[0].font.size = Pt(10)
        doc.add_paragraph()

    subject_line = f"{report.subject_address or ''}, {report.subject_city or ''}".strip(", ") or "описания имот"
    _intro_block(
        "Предмет на заданието",
        f"Оценка на пазарната стойност на недвижим имот, находящ се на {subject_line}"
        + (f" (кадастрален идентификатор {report.subject_cadastral_id})" if report.subject_cadastral_id else "")
        + ".",
    )
    _intro_block("Цел на оценката", purpose_texts["cel"])
    _intro_block("Стандарт и база на стойността", purpose_texts["standart"])

    appraiser_name = (appraiser.full_name or appraiser.username) if appraiser else None
    appraiser_cert = appraiser.appraiser_certificate_no if appraiser else None
    p_decl_title = doc.add_paragraph()
    r_decl_title = p_decl_title.add_run("Декларация за независимост")
    r_decl_title.bold = True
    r_decl_title.font.size = Pt(10.5)
    for line in [
        "не е свързано лице с възложителя или собственика на оценявания имот;",
        "няма настоящ или бъдещ имуществен интерес, свързан с обекта на оценката;",
        "възнаграждението за изготвяне на доклада не е обвързано с заключената стойност.",
    ]:
        p_line = doc.add_paragraph(f"•  {line}", style=None)
        p_line.runs[0].font.size = Pt(10)
    if appraiser_name:
        p_app = doc.add_paragraph()
        r_app_lbl = p_app.add_run("Изготвил оценката: ")
        r_app_lbl.bold = True
        r_app_lbl.font.size = Pt(10)
        cert_suffix = f" (сертификат № {appraiser_cert})" if appraiser_cert else ""
        r_app_val = p_app.add_run(f"{appraiser_name}{cert_suffix}")
        r_app_val.font.size = Pt(10)

    doc.add_page_break()

    section_num = 1

    # ── Section: Subject property ───────────────────────────────
    h1 = doc.add_heading(f"{section_num}. Описание на оценявания имот", level=1)
    if h1.runs:
        h1.runs[0].font.color.rgb = _BRAND_DARK

    tbl_subj = doc.add_table(rows=1, cols=2)
    tbl_subj.style = "Table Grid"
    tbl_subj.autofit = False
    tbl_subj._tbl.remove(tbl_subj.rows[0]._tr)

    def _subj_row(label, value):
        row = tbl_subj.add_row()
        _fill_cell(row.cells[0], label, bold=True, pt=10)
        _fill_cell(row.cells[1], value if value else "—", pt=10)

    _subj_row("Адрес", report.subject_address)
    _subj_row("Град", report.subject_city)
    _subj_row("Площ (кв.м)", _fmt(report.subject_area_sqm))
    _subj_row(
        "Етаж / Общо етажи",
        f"{report.subject_floor or '—'} / {report.subject_total_floors or '—'}",
    )
    _subj_row("Тип строителство", report.subject_construction)
    _subj_row("Година на строеж", str(report.subject_year) if report.subject_year else None)
    _subj_row(
        "Дата на оценка",
        report.valuation_date.strftime("%d.%m.%Y") if report.valuation_date else None,
    )

    if report.subject_description:
        doc.add_paragraph()
        pd = doc.add_paragraph()
        r_lbl = pd.add_run("Описание: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(10)
        r_desc = pd.add_run(report.subject_description)
        r_desc.font.size = Pt(10)

    doc.add_paragraph()

    # ── Section: Legal & zoning status (conditional on GIS data present) ──
    if report.legal_description:
        section_num += 1
        h_legal = doc.add_heading(f"{section_num}. Правно и градоустройствено състояние", level=1)
        if h_legal.runs:
            h_legal.runs[0].font.color.rgb = _BRAND_DARK
        p_legal = doc.add_paragraph(report.legal_description)
        p_legal.runs[0].font.size = Pt(10)
        if report.subject_cadastral_id:
            p_cad = doc.add_paragraph()
            r_cad_lbl = p_cad.add_run("Кадастрален идентификатор: ")
            r_cad_lbl.bold = True
            r_cad_lbl.font.size = Pt(9)
            r_cad_val = p_cad.add_run(report.subject_cadastral_id)
            r_cad_val.font.size = Pt(9)
        doc.add_paragraph()

    # ── Section: Sales comparables ──────────────────────────────
    section_num += 1
    h2 = doc.add_heading(f"{section_num}. Пазарен подход — продажни сравними", level=1)
    if h2.runs:
        h2.runs[0].font.color.rgb = _BRAND_DARK

    market_ctx = _market_context_paragraph(db, report)
    if market_ctx:
        p_ctx = doc.add_paragraph(market_ctx)
        p_ctx.runs[0].italic = True
        p_ctx.runs[0].font.size = Pt(9)
        p_ctx.runs[0].font.color.rgb = _MUTED
        doc.add_paragraph()

    if report.submarket_rationale:
        p_sr_lbl = doc.add_paragraph()
        r_sr_lbl = p_sr_lbl.add_run("Обосновка на съпоставимата зона: ")
        r_sr_lbl.bold = True
        r_sr_lbl.font.size = Pt(9.5)
        r_sr_txt = p_sr_lbl.add_run(report.submarket_rationale)
        r_sr_txt.font.size = Pt(9.5)
        doc.add_paragraph()

    if pinned_sale:
        _write_comp_table(doc, pinned_sale, "sale", report, pool_sale["stats"])
        _write_adjustment_breakdown_table(doc, pinned_sale)
    else:
        p_no = doc.add_paragraph("Няма закачени (\U0001F4CC) продажни сравними за доклада.")
        p_no.runs[0].italic = True
        p_no.runs[0].font.size = Pt(10)

    # Concluded sales value
    if report.concluded_value_sales:
        doc.add_paragraph()
        p_cs = doc.add_paragraph()
        r_cs = p_cs.add_run(
            f"Стойност по пазарен подход: "
            f"{_fmt(report.concluded_value_sales)} {report.concluded_currency or 'EUR'}"
        )
        r_cs.bold = True
        r_cs.font.size = Pt(11)
        r_cs.font.color.rgb = _BRAND_DARK

    # ── Section: Rent comparables (if any) ──────────────────────
    if pinned_rent or pool_rent["total_count"] > 0:
        section_num += 1
        doc.add_paragraph()
        h3 = doc.add_heading(f"{section_num}. Доходен подход — наемни сравними", level=1)
        if h3.runs:
            h3.runs[0].font.color.rgb = _BRAND_DARK

        if pinned_rent:
            _write_comp_table(doc, pinned_rent, "rent", report, pool_rent["stats"])
            _write_adjustment_breakdown_table(doc, pinned_rent)
        else:
            p_no2 = doc.add_paragraph("Няма закачени (\U0001F4CC) наемни сравними за доклада.")
            p_no2.runs[0].italic = True
            p_no2.runs[0].font.size = Pt(10)

        if report.concluded_value_income:
            doc.add_paragraph()
            p_ci = doc.add_paragraph()
            r_ci = p_ci.add_run(
                f"Стойност по доходен подход: "
                f"{_fmt(report.concluded_value_income)} {report.concluded_currency or 'EUR'}"
            )
            r_ci.bold = True
            r_ci.font.size = Pt(11)
            r_ci.font.color.rgb = _BRAND_DARK

    # ── Section: Concluded value ────────────────────────
    section_num += 1
    conc_num = section_num
    doc.add_paragraph()
    h_conc = doc.add_heading(f"{conc_num}. Заключение — оценена стойност", level=1)
    if h_conc.runs:
        h_conc.runs[0].font.color.rgb = _BRAND_DARK

    tbl_conc = doc.add_table(rows=1, cols=3)
    tbl_conc.style = "Table Grid"
    tbl_conc._tbl.remove(tbl_conc.rows[0]._tr)

    def _weight_str(w) -> str:
        return f"{_fmt(w, 0)}%" if w is not None else "—"

    if report.concluded_value_sales:
        r = tbl_conc.add_row()
        _fill_cell(r.cells[0], "Пазарна стойност (пазарен подход)", bold=True, pt=10)
        _fill_cell(
            r.cells[1],
            f"{_fmt(report.concluded_value_sales)} {report.concluded_currency or 'EUR'}",
            pt=10,
        )
        _fill_cell(r.cells[2], _weight_str(report.weight_sales_pct), pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    if report.concluded_value_income:
        r = tbl_conc.add_row()
        _fill_cell(r.cells[0], "Пазарна стойност (доходен подход)", bold=True, pt=10)
        _fill_cell(
            r.cells[1],
            f"{_fmt(report.concluded_value_income)} {report.concluded_currency or 'EUR'}",
            pt=10,
        )
        _fill_cell(r.cells[2], _weight_str(report.weight_income_pct), pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    if report.concluded_value_residual:
        r = tbl_conc.add_row()
        _fill_cell(r.cells[0], "Стойност на парцела (остатъчен метод)", bold=True, pt=10)
        _fill_cell(
            r.cells[1],
            f"{_fmt(report.concluded_value_residual)} {report.concluded_currency or 'EUR'}",
            pt=10,
        )
        _fill_cell(r.cells[2], _weight_str(report.weight_residual_pct), pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    if report.concluded_value:
        r = tbl_conc.add_row()
        _fill_cell(
            r.cells[0], "КРАЙНА ОЦЕНЕНА СТОЙНОСТ (претеглена)",
            bold=True, pt=11, color=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _fill_cell(
            r.cells[1],
            f"{_fmt(report.concluded_value)} {report.concluded_currency or 'EUR'}",
            bold=True, pt=11, color=_WHITE,
        )
        _fill_cell(r.cells[2], "", pt=11, color=_WHITE)
        _set_cell_bg(r.cells[0], "1E3A5F")
        _set_cell_bg(r.cells[1], "1E3A5F")
        _set_cell_bg(r.cells[2], "1E3A5F")
    elif not (report.concluded_value_sales or report.concluded_value_income):
        p_nv = doc.add_paragraph("Крайната оценена стойност не е въведена.")
        p_nv.runs[0].italic = True
        p_nv.runs[0].font.size = Pt(10)

    if report.weighting_rationale:
        doc.add_paragraph()
        p_wr = doc.add_paragraph()
        r_wr_lbl = p_wr.add_run("Обосновка на теглата: ")
        r_wr_lbl.bold = True
        r_wr_lbl.font.size = Pt(10)
        r_wr_txt = p_wr.add_run(report.weighting_rationale)
        r_wr_txt.font.size = Pt(10)

    if report.appraiser_notes:
        doc.add_paragraph()
        p_an = doc.add_paragraph()
        r_an_lbl = p_an.add_run("Бележки на оценителя: ")
        r_an_lbl.bold = True
        r_an_lbl.font.size = Pt(10)
        r_an_txt = p_an.add_run(report.appraiser_notes)
        r_an_txt.font.size = Pt(10)

    # ── КНОБ statement in document body ──────────────────────────
    doc.add_paragraph()
    p_knob = doc.add_paragraph(_KNOB_TEXT)
    p_knob.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_knob.runs[0].font.size = Pt(8.5)
    p_knob.runs[0].font.color.rgb = _MUTED

    # ── Appendix: Limiting conditions (generic, purpose-independent) ──────
    doc.add_page_break()
    h_lc = doc.add_heading("Приложение — Ограничаващи условия и допускания", level=1)
    if h_lc.runs:
        h_lc.runs[0].font.color.rgb = _BRAND_DARK
    p_lc = doc.add_paragraph(_LIMITING_CONDITIONS_TEXT)
    p_lc.runs[0].font.size = Pt(9.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Draft / report management ─────────────────────────────────────────────────

def get_or_create_draft(db: Session, user_id: int) -> AppraisalReport:
    draft = (
        db.query(AppraisalReport)
        .filter(AppraisalReport.status == "draft", AppraisalReport.user_id == user_id)
        .order_by(AppraisalReport.updated_at.desc())
        .first()
    )
    if draft is None:
        draft = _new_draft_obj(db, user_id)
    return draft


def new_draft(db: Session, user_id: int) -> AppraisalReport:
    return _new_draft_obj(db, user_id)


def _new_draft_obj(db: Session, user_id: int) -> AppraisalReport:
    draft = AppraisalReport(title="Нов доклад", status="draft", user_id=user_id)
    db.add(draft)
    db.commit()
    db.refresh(draft)
    return draft


def get_user_reports(db: Session, user_id: int) -> list[dict]:
    rows = db.execute(text("""
        SELECT
            ar.id,
            ar.title,
            ar.status,
            ar.created_at,
            ar.updated_at,
            ar.subject_address,
            ar.subject_city,
            ar.valuation_date,
            count(cp.id) FILTER (WHERE cp.comparable_type = 'sale') AS sale_count,
            count(cp.id) FILTER (WHERE cp.comparable_type = 'rent') AS rent_count,
            count(cp.id) FILTER (WHERE cp.pinned_for_report = true)  AS pinned_count
        FROM appraisal_reports ar
        LEFT JOIN comparable_pool cp ON cp.report_id = ar.id
        WHERE ar.user_id = :uid
        GROUP BY ar.id
        ORDER BY ar.updated_at DESC
    """), {"uid": user_id}).mappings().all()
    return [dict(r) for r in rows]


def get_report_for_user(
    db: Session, report_id: uuid.UUID, user_id: int
) -> AppraisalReport | None:
    return (
        db.query(AppraisalReport)
        .filter(AppraisalReport.id == report_id, AppraisalReport.user_id == user_id)
        .first()
    )


def update_subject(db: Session, report_id: uuid.UUID, data: dict) -> None:
    report = db.get(AppraisalReport, report_id)
    if not report:
        return
    fields = (
        "title", "subject_address", "subject_city", "subject_area_sqm",
        "subject_floor", "subject_total_floors", "subject_construction",
        "subject_year", "subject_description", "valuation_date",
        "subject_property_type", "subject_geo_category", "subject_neighborhood",
        "subject_cadastral_id",
    )
    for f in fields:
        if f in data:
            setattr(report, f, data[f] if data[f] != "" else None)
    # report_purpose is NOT NULL -- unlike the fields above, an empty
    # submitted value must be ignored (keep the existing/default value)
    # rather than nulled out.
    if data.get("report_purpose"):
        report.report_purpose = data["report_purpose"]
    db.commit()


def update_income_approach(
    db: Session,
    report_id: uuid.UUID,
    rent_per_sqm_month: float | None,
    cap_rate_pct: float | None,
    concluded_per_sqm: float | None,
    subject_area_sqm: float | None,
) -> None:
    report = db.get(AppraisalReport, report_id)
    if not report:
        return
    if rent_per_sqm_month is not None:
        report.annual_rent_estimate = round(rent_per_sqm_month * 12, 2)
    if cap_rate_pct is not None:
        report.capitalization_rate = round(cap_rate_pct / 100, 6)
    if concluded_per_sqm is not None:
        area = subject_area_sqm or float(report.subject_area_sqm or 0)
        report.concluded_value_income = round(concluded_per_sqm * area, 2) if area > 0 else round(concluded_per_sqm, 2)
    db.commit()


def update_sales_approach(
    db: Session,
    report_id: uuid.UUID,
    concluded_value_sales: float | None,
    source: str = "manual",
) -> None:
    report = db.get(AppraisalReport, report_id)
    if not report:
        return
    if concluded_value_sales is not None:
        report.concluded_value_sales = round(concluded_value_sales, 2)
        report.concluded_value_sales_source = source
    db.commit()


def update_legal_description(
    db: Session,
    report_id: uuid.UUID,
    text: str | None,
    source: str = "manual",
) -> None:
    report = db.get(AppraisalReport, report_id)
    if not report:
        return
    report.legal_description = text or None
    report.legal_description_source = source if text else None
    db.commit()


def update_submarket_rationale(db: Session, report_id: uuid.UUID, text: str | None) -> None:
    report = db.get(AppraisalReport, report_id)
    if not report:
        return
    report.submarket_rationale = text or None
    db.commit()


def update_income_market_rationale(db: Session, report_id: uuid.UUID, text: str | None) -> None:
    report = db.get(AppraisalReport, report_id)
    if not report:
        return
    report.income_market_rationale = text or None
    db.commit()


def update_residual_approach(
    db: Session,
    report_id: uuid.UUID,
    concluded_value_residual: float | None,
) -> None:
    report = db.get(AppraisalReport, report_id)
    if not report:
        return
    if concluded_value_residual is not None:
        report.concluded_value_residual = round(concluded_value_residual, 2)
    db.commit()


# Bounds + defaults for compute_income_valuation()'s assumption params
# (Phase 7, Tier 5 -- app/services/llm/tools.py's compute_income_valuation
# tool enforces these server-side regardless of what the model requests, so
# an out-of-range value from the model is clamped, never silently trusted).
# Defaults match _income_analysis.html's JS panel's own pre-filled values,
# so a report with no manually-saved income inputs gets the same starting
# assumptions either way.
INCOME_ASSUMPTION_BOUNDS = {
    "expenses_pct": (10.0, 35.0),
    "vacancy_pct": (2.0, 15.0),
    "cap_rate_pct": (4.0, 12.0),
    "growth_pct": (-2.0, 6.0),
    "period_years": (3, 10),
    "terminal_cap_rate_pct": (5.0, 12.0),
}
INCOME_ASSUMPTION_DEFAULTS = {
    "expenses_pct": 20.0,
    "vacancy_pct": 8.0,
    "cap_rate_pct": 7.0,
    "growth_pct": 2.0,
    "period_years": 5,
    "terminal_cap_rate_pct": 7.5,
}


def compute_income_valuation(
    rent_per_sqm_month: float,
    sale_price_per_sqm: float | None,
    expenses_pct: float,
    vacancy_pct: float,
    cap_rate_pct: float,
    growth_pct: float,
    period_years: int,
    terminal_cap_rate_pct: float,
) -> dict:
    """Direct capitalization + multi-year DCF with terminal value.

    Pure Python port of _income_analysis.html's calcIncome() JS -- kept in
    exact numeric parity (same formula, same variable roles) so this
    produces the same figures the manual UI panel would show for identical
    inputs. NOI = annual rent x (1-expenses) x (1-vacancy); the DCF
    discounts NOI at cap_rate_pct (used as the discount rate, matching the
    existing UI's own convention -- "use cap rate as discount rate") over
    period_years with growth_pct annual rent growth, plus a terminal value
    (NOI at year period+1 / terminal_cap_rate_pct, discounted back at the
    same rate). All percentages are plain numbers (7.0 means 7%, not 0.07).
    """
    expenses = expenses_pct / 100
    vacancy = vacancy_pct / 100
    cap_rate = cap_rate_pct / 100
    growth = growth_pct / 100
    terminal_cap_rate = terminal_cap_rate_pct / 100

    annual_rent = rent_per_sqm_month * 12
    effective_pct = (1 - expenses) * (1 - vacancy)
    noi = annual_rent * effective_pct

    gross_yield_pct = (annual_rent / sale_price_per_sqm * 100) if sale_price_per_sqm else None
    net_yield_pct = (noi / sale_price_per_sqm * 100) if sale_price_per_sqm else None
    direct_value = (noi / cap_rate) if cap_rate > 0 else None

    rows = []
    pv = 0.0
    current_noi = noi
    r = cap_rate
    for t in range(1, int(period_years) + 1):
        pv_factor = 1 / ((1 + r) ** t)
        pv_noi = current_noi * pv_factor
        rows.append({
            "year": t,
            "noi": round(current_noi, 2),
            "pv_factor": round(pv_factor, 4),
            "pv_noi": round(pv_noi, 2),
        })
        pv += pv_noi
        current_noi *= (1 + growth)

    noi_terminal = noi * ((1 + growth) ** int(period_years))
    tv_undiscounted = (noi_terminal / terminal_cap_rate) if terminal_cap_rate > 0 else 0.0
    pv_tv = tv_undiscounted / ((1 + r) ** int(period_years))
    dcf_value = pv + pv_tv

    return {
        "gross_yield_pct": round(gross_yield_pct, 2) if gross_yield_pct is not None else None,
        "net_yield_pct": round(net_yield_pct, 2) if net_yield_pct is not None else None,
        "noi_per_sqm_year": round(noi, 2),
        "direct_value_per_sqm": round(direct_value, 2) if direct_value is not None else None,
        "dcf_value_per_sqm": round(dcf_value, 2),
        "dcf_rows": rows,
        "terminal_value_pv_per_sqm": round(pv_tv, 2),
        "terminal_value_undiscounted_per_sqm": round(tv_undiscounted, 2),
    }


def compute_weighted_conclusion(
    concluded_value_sales: float | None,
    concluded_value_income: float | None,
    concluded_value_residual: float | None,
    weight_sales_pct: float | None,
    weight_income_pct: float | None,
    weight_residual_pct: float | None,
) -> float | None:
    """Normalized weighted average of whichever approaches have BOTH a
    saved value and a positive weight -- normalized (divided by the sum of
    weights actually used) so the weights don't need to add up to exactly
    100 for a sane result, matching how an appraiser might leave one
    approach's weight blank rather than force the other two to compensate."""
    pairs = [
        (concluded_value_sales, weight_sales_pct),
        (concluded_value_income, weight_income_pct),
        (concluded_value_residual, weight_residual_pct),
    ]
    used = [(v, w) for v, w in pairs if v is not None and w is not None and w > 0]
    if not used:
        return None
    total_weight = sum(w for _, w in used)
    return round(sum(v * w for v, w in used) / total_weight, 2)


def update_conclusion(
    db: Session,
    report_id: uuid.UUID,
    weight_sales_pct: float | None,
    weight_income_pct: float | None,
    weight_residual_pct: float | None,
    weighting_rationale: str | None,
) -> None:
    report = db.get(AppraisalReport, report_id)
    if not report:
        return
    report.weight_sales_pct = weight_sales_pct
    report.weight_income_pct = weight_income_pct
    report.weight_residual_pct = weight_residual_pct
    report.weighting_rationale = weighting_rationale or None
    report.concluded_value = compute_weighted_conclusion(
        float(report.concluded_value_sales) if report.concluded_value_sales is not None else None,
        float(report.concluded_value_income) if report.concluded_value_income is not None else None,
        float(report.concluded_value_residual) if report.concluded_value_residual is not None else None,
        weight_sales_pct, weight_income_pct, weight_residual_pct,
    )
    db.commit()


def delete_user_report(db: Session, report_id: uuid.UUID, user_id: int) -> None:
    report = (
        db.query(AppraisalReport)
        .filter(AppraisalReport.id == report_id, AppraisalReport.user_id == user_id)
        .first()
    )
    if report:
        db.delete(report)
        db.commit()


def finalize_user_report(db: Session, report_id: uuid.UUID, user_id: int) -> None:
    report = (
        db.query(AppraisalReport)
        .filter(AppraisalReport.id == report_id, AppraisalReport.user_id == user_id)
        .first()
    )
    if report and report.status == "draft":
        report.status = "finalized"
        db.commit()


# ── Pool operations ───────────────────────────────────────────────────────────

def add_to_pool(
    db: Session,
    listing_ids: list[int],
    comparable_type: str,
    report_id: uuid.UUID,
    user_id: int,
) -> int:
    if not listing_ids:
        return 0
    stmt = (
        pg_insert(ComparablePool)
        .values([{
            "listing_id": lid,
            "comparable_type": comparable_type,
            "report_id": report_id,
            "user_id": user_id,
        } for lid in listing_ids])
        .on_conflict_do_nothing(constraint="uq_pool_listing_ctype_report")
    )
    result = db.execute(stmt)
    db.commit()
    return result.rowcount


def remove_from_pool(db: Session, pool_id: int) -> None:
    item = db.get(ComparablePool, pool_id)
    if item:
        db.delete(item)
        db.commit()


def clear_pool(
    db: Session, report_id: uuid.UUID, comparable_type: str | None = None
) -> None:
    q = db.query(ComparablePool).filter(ComparablePool.report_id == report_id)
    if comparable_type:
        q = q.filter_by(comparable_type=comparable_type)
    q.delete()
    db.commit()


def toggle_pin(db: Session, pool_id: int) -> bool:
    """Toggle pinned_for_report. Returns new value. Enforces MAX_PINNED."""
    item = db.get(ComparablePool, pool_id)
    if not item:
        return False
    if not item.pinned_for_report:
        pinned_count = (
            db.query(ComparablePool)
            .filter_by(
                comparable_type=item.comparable_type,
                report_id=item.report_id,
                pinned_for_report=True,
            )
            .count()
        )
        if pinned_count >= MAX_PINNED:
            return False
    item.pinned_for_report = not item.pinned_for_report
    db.commit()
    return item.pinned_for_report


ADJUSTMENT_FACTOR_LABELS: dict[str, str] = {
    "market":    "Пазарни условия",
    "location":  "Местоположение",
    "size":      "Площ",
    "floor":     "Етаж",
    "condition": "Строителство / състояние",
}


def update_pool_adjustment(
    db: Session,
    pool_id: int,
    adjustment_pct: float | None,
    analyst_note: str,
    adjustment_factors: dict[str, float] | None = None,
) -> None:
    """When adjustment_factors is given (even an empty dict clears factor
    mode), adjustment_pct is DERIVED as their sum and the explicit
    adjustment_pct argument is ignored -- keeps the two representations from
    silently disagreeing. Passing adjustment_factors=None (the default)
    preserves the older single-blended-% entry mode untouched."""
    item = db.get(ComparablePool, pool_id)
    if not item:
        return
    if adjustment_factors is not None:
        cleaned = {k: v for k, v in adjustment_factors.items() if v}
        item.adjustment_factors = cleaned or None
        item.adjustment_pct = round(sum(cleaned.values()), 2) if cleaned else None
    else:
        item.adjustment_pct = adjustment_pct
    item.analyst_note = analyst_note or None
    db.commit()


# ── Pool query + statistics ───────────────────────────────────────────────────

def get_pool_with_stats(
    db: Session, comparable_type: str, report_id: uuid.UUID
) -> dict:
    """Returns {rows, stats, pinned_count, total_count} for the given report + type."""
    rows = db.execute(text("""
        SELECT
            cp.id AS pool_id,
            cp.pinned_for_report,
            cp.adjustment_pct,
            cp.adjustment_factors,
            cp.analyst_note,
            cp.added_at,
            l.id AS listing_id,
            l.ad_url,
            l.title,
            l.title_city_model,
            l.title_geo_2_model,
            l.location_raw,
            l.total_price,
            l.currency,
            l.price_per_sqm_model,
            l.area_sqm_model,
            l.floor_model,
            l.total_floors_model,
            l.construction_type_model,
            l.construction_year_model,
            l.published_date,
            l.last_seen_at,
            l.deal_type_normalized,
            l.property_type_raw
        FROM comparable_pool cp
        JOIN listings l ON l.id = cp.listing_id
        WHERE cp.comparable_type = :ctype
          AND cp.report_id = :rid
          AND l.status = 'active'
        ORDER BY cp.pinned_for_report DESC, cp.added_at DESC
    """), {"ctype": comparable_type, "rid": str(report_id)}).mappings().all()

    items = []
    for r in rows:
        r = dict(r)
        adj = float(r["adjustment_pct"]) if r["adjustment_pct"] is not None else 0.0
        ppsqm = float(r["price_per_sqm_model"]) if r["price_per_sqm_model"] else None
        r["adj_ppsqm"] = round(ppsqm * (1 + adj / 100), 0) if ppsqm is not None else None
        r["adjustment_pct"] = adj
        items.append(r)

    stats = _compute_stats(db, comparable_type, report_id)
    pinned_count = sum(1 for i in items if i["pinned_for_report"])
    return {
        "rows": items,
        "stats": stats,
        "pinned_count": pinned_count,
        "total_count": len(items),
    }


def _compute_stats(
    db: Session, comparable_type: str, report_id: uuid.UUID
) -> dict | None:
    row = db.execute(text("""
        SELECT
            count(*) AS n,
            round(min(l.price_per_sqm_model)::numeric, 0)  AS min_ppsqm,
            round(max(l.price_per_sqm_model)::numeric, 0)  AS max_ppsqm,
            round(avg(l.price_per_sqm_model)::numeric, 0)  AS mean_ppsqm,
            round(percentile_cont(0.25) WITHIN GROUP (ORDER BY l.price_per_sqm_model)::numeric, 0) AS p25,
            round(percentile_cont(0.50) WITHIN GROUP (ORDER BY l.price_per_sqm_model)::numeric, 0) AS median,
            round(percentile_cont(0.75) WITHIN GROUP (ORDER BY l.price_per_sqm_model)::numeric, 0) AS p75,
            round(min(l.area_sqm_model)::numeric, 0)       AS min_area,
            round(max(l.area_sqm_model)::numeric, 0)       AS max_area,
            round(avg(l.area_sqm_model)::numeric, 0)       AS mean_area,
            round(min(l.total_price)::numeric, 0)          AS min_price,
            round(max(l.total_price)::numeric, 0)          AS max_price,
            round(avg(l.total_price)::numeric, 0)          AS mean_price
        FROM comparable_pool cp
        JOIN listings l ON l.id = cp.listing_id
        WHERE cp.comparable_type = :ctype
          AND cp.report_id = :rid
          AND l.status = 'active'
          AND l.price_per_sqm_model IS NOT NULL
    """), {"ctype": comparable_type, "rid": str(report_id)}).mappings().first()

    if not row or not row["n"]:
        return None
    return dict(row)


# ── Excel export ──────────────────────────────────────────────────────────────

_HEADER_FONT = Font(bold=True, color="FFFFFF")
_HEADER_FILL = PatternFill("solid", fgColor="2563EB")
_PINNED_FILL = PatternFill("solid", fgColor="DCFCE7")
_SUBJECT_FILL = PatternFill("solid", fgColor="FEF3C7")
_CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)


def _autowidth(ws) -> None:
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 3, 45)


def _write_header(ws, cols: list[str]) -> None:
    ws.append(cols)
    for cell in ws[1]:
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _CENTER


def _floor_str(item: dict) -> str:
    f, t = item.get("floor_model"), item.get("total_floors_model")
    if f is None and t is None:
        return "—"
    return f"{f if f is not None else '?'}/{t}" if t else (str(f) if f is not None else "—")


def export_excel(db: Session, report: AppraisalReport) -> io.BytesIO:
    pool_sale = get_pool_with_stats(db, "sale", report.id)
    pool_rent = get_pool_with_stats(db, "rent", report.id)

    wb = Workbook()
    wb.remove(wb.active)

    def _subject_row(ws, ncols: int) -> None:
        if not (report and report.subject_address):
            return
        pad = [""] * ncols
        row_data = [
            "Обект",
            f"{report.subject_address or ''}\n{report.subject_city or ''}".strip(),
            float(report.subject_area_sqm) if report.subject_area_sqm else "",
            "", "",
            report.subject_construction or "",
            report.subject_year or "",
            f"{report.subject_floor or '?'}/{report.subject_total_floors or '?'}",
        ] + pad[8:]
        ws.append(row_data[:ncols])
        for cell in ws[ws.max_row]:
            cell.fill = _SUBJECT_FILL

    def _stats_row(ws, stats: dict | None, ncols: int) -> None:
        if not stats:
            return
        row = [
            "СТАТИСТИКИ", f"N={stats['n']}",
            f"Площ: {stats['min_area']}–{stats['max_area']} (ср.{stats['mean_area']})",
            "", "",
            f"Цена/кв.м: мин {stats['min_ppsqm']} | ср {stats['mean_ppsqm']} | макс {stats['max_ppsqm']}",
            f"Q25={stats['p25']} | Медиана={stats['median']} | Q75={stats['p75']}",
        ]
        ws.append((row + [""] * ncols)[:ncols])
        fill = PatternFill("solid", fgColor="EFF6FF")
        for cell in ws[ws.max_row]:
            cell.fill = fill

    def _write_sheet(ws, items, stats, cols, include_pin_col=True):
        all_cols = (["📌"] if include_pin_col else []) + cols
        _write_header(ws, all_cols)
        _subject_row(ws, len(all_cols))
        _stats_row(ws, stats, len(all_cols))
        for pos, item in enumerate(items, start=1):
            row = ([("✔" if item["pinned_for_report"] else "")] if include_pin_col else []) + [
                pos,
                f"{item.get('title_city_model') or ''} {item.get('title_geo_2_model') or ''}\n"
                f"{item.get('location_raw') or ''}".strip(),
                float(item["area_sqm_model"]) if item.get("area_sqm_model") else "",
                float(item["total_price"]) if item.get("total_price") else "",
                float(item["price_per_sqm_model"]) if item.get("price_per_sqm_model") else "",
                item.get("construction_type_model") or "",
                item.get("construction_year_model") or "",
                _floor_str(item),
                item["adjustment_pct"],
                item["adj_ppsqm"],
                item.get("analyst_note") or "",
                item["ad_url"],
            ]
            ws.append(row[:len(all_cols)])
            if item["pinned_for_report"]:
                for cell in ws[ws.max_row]:
                    cell.fill = _PINNED_FILL
        _autowidth(ws)

    sale_cols = ["№", "Адрес / Описание", "Площ (кв.м)", "Цена (EUR)",
                 "Цена/кв.м (EUR)", "Строит. тип", "Год.", "Ет./Обш.",
                 "Корекция (%)", "Кор. цена/кв.м", "Бележки", "URL"]
    rent_cols = ["№", "Адрес / Описание", "Площ (кв.м)", "Наем/мес (EUR)",
                 "Наем/кв.м/мес (EUR)", "Строит. тип", "Год.", "Ет./Обш.",
                 "Корекция (%)", "Кор. наем/кв.м", "Бележки", "URL"]

    if pool_sale["rows"]:
        ws = wb.create_sheet("Продажни сравними")
        _write_sheet(ws, pool_sale["rows"], pool_sale["stats"], sale_cols)

    if pool_rent["rows"]:
        ws = wb.create_sheet("Наемни сравними")
        _write_sheet(ws, pool_rent["rows"], pool_rent["stats"], rent_cols)

    if not wb.sheetnames:
        ws = wb.create_sheet("Без данни")
        ws.append(["Няма добавени сравними."])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
